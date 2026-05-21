import streamlit as st
import google.generativeai as genai
from datetime import datetime
from io import BytesIO
from docx import Document

# Configuração da página - Deve ser a primeira coisa no Streamlit
st.set_page_config(page_title="Assistente de Prontuário", layout="centered")

st.title("📝 Assistente de Prontuário Clínico")

# --- FUNÇÃO PARA GERAR O ARQUIVO .DOCX ---
def gerar_docx(texto_prontuario):
    doc = Document()
    
    # Divide o texto por quebras de linha para manter a estrutura e formatação no Word
    linhas = texto_prontuario.split('\n')
    for linha in linhas:
        linha_limpa = linha.strip()
        if (linha_limpa.startswith("ANOTAÇÕES CLÍNICAS") or 
            linha_limpa.startswith("Cliente:") or 
            linha_limpa.startswith("Data:") or 
            linha_limpa.startswith("Desenvolvimento da Sessão:") or 
            linha_limpa.startswith("Manejo e Intervenções:") or 
            linha_limpa.startswith("Pontos Relevantes para a Continuidade:")):
            
            p = doc.add_paragraph()
            p.add_run(linha).bold = True
        else:
            doc.add_paragraph(linha)
            
    conteudo_docx = BytesIO()
    doc.save(conteudo_docx)
    conteudo_docx.seek(0)
    return conteudo_docx

# --- GERENCIAMENTO SEGURO DA API KEY ---
try:
    gemini_key = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=gemini_key)
    model = genai.GenerativeModel('gemini-3.1-flash-lite')
    api_pronta = True
except Exception as e:
    st.error("Erro na configuração da API Key. Verifique as configurações de Secrets.")
    api_pronta = False

# --- LÓGICA DA APLICAÇÃO ---
if api_pronta:
    # Seleção de Abordagem Teórica
    abordagem = st.selectbox(
        "Selecione a abordagem teórica para a redação:",
        ["Análise do Comportamento", "Histórico-Cultural"]
    )
    
    paciente_ref = st.text_input("Nome / Iniciais / Referência do Paciente", placeholder="Ex: F.B./Paciente A")
    
    placeholder_notas = (
        "Como o(a) paciente chegou e qual o tema principal de hoje?\n"
        "Quais manejos e intervenções você utilizou?\n"
        "O que foi observado ao final e o que se espera para as próximas sessões?"
    )
    
    anotacoes = st.text_area(
        "Digite as notas da sessão aqui:", 
        height=250, 
        placeholder=placeholder_notas
    )

    if anotacoes:
        if st.button("Gerar Prontuário Estruturado", key="btn_gerar"):
            with st.spinner("Carregando..."):
                try:
                    data_atual = datetime.now().strftime("%d/%m/%Y")
                    nome_cliente = paciente_ref.strip() if paciente_ref else "Não informado"

                    # --- DIRETRIZES TÉCNICAS ESPECÍFICAS POR ABORDAGEM ---
                    if abordagem == "Análise do Comportamento":
                        diretriz_abordagem = (
                            "- Orientação Teórica (Análise do Comportamento): Redija priorizando conceitos como "
                            "contingências de reforçamento, funções do comportamento, variáveis antecedentes e consequentes, "
                            "repertório comportamental e regras/autorregras. Evite termos mentalistas ou internalistas."
                        )
                    else:  # Histórico-Cultural
                        diretriz_abordagem = (
                            "- Orientação Teórica (Psicologia Histórico-Cultural): Redija priorizando conceitos como "
                            "mediação semiótica, funções psicológicas superiores, vivência (perejivanie), constituição social "
                            "do psiquismo e dimensões histórico-sociais que atravessam a demanda do sujeito."
                        )

                    prompt_completo = f"""
                    Atue como um Assistente de Escrita Clínico para Psicólogos. Sua tarefa é converter as notas de sessão fornecidas pelo profissional em um prontuário técnico, ético e extremamente enxuto, seguindo rigorosamente a estrutura abaixo.
                    
                    Você DEVE obrigatoriamente iniciar o seu output com o cabeçalho exatamente no formato estruturado abaixo. É CRUCIAL que você mantenha cada informação em uma linha separada, respeitando estritamente as quebras de linha. NÃO junte estas linhas em um único parágrafo.
                    
                    
                    ANOTAÇÕES CLÍNICAS DE ATENDIMENTO
                    Cliente: {nome_cliente}
                    Data: {data_atual}
                    
                    
                    Diretrizes de Escrita:
                    - Tom Profissional: Use uma linguagem técnica e estritamente adequada à abordagem selecionada abaixo.
                    {diretriz_abordagem}
                    - Neutralidade: Não invente diagnósticos ou sentimentos que não foram descritos.
                    - Concisão: Mantenha cada parágrafo curto. Se uma informação não estiver nas notas, use termos genéricos como "não reportado" ou apenas foque no que existe.
                    - Flexibilidade de Gênero: Adapte o gênero (o/a cliente) conforme as notas, ou use termos neutros como "Paciente".
                    
                    Estrutura do Restante do Output (pule uma linha após o cabeçalho):
                    
                    Desenvolvimento da Sessão:
                    [Descreva o estado inicial/postura se houver na nota, ou inicie diretamente com o tema trazido. Relate o conteúdo central trabalhado de forma resumida e clinicamente relevante sob a ótica da abordagem epistemológica escolhida].
                    
                    Manejo e Intervenções:
                    [Descreva a postura técnica adotada (escuta, acolhimento, análise funcional ou mediação conforme a teoria) e quais intervenções foram realizadas com base no relato].
                    
                    Pontos Relevantes para a Continuidade:
                    [Registre padrões observados, dificuldades ou combinados para o futuro].
                    
                    Notas para processar: {anotacoes}
                    """
                    
                    response = model.generate_content(prompt_completo)
                    
                    st.session_state["texto_gerado"] = response.text
                    st.session_state["nome_arquivo"] = nome_cliente.replace(" ", "_")
                    st.session_state["data_arquivo"] = data_atual.replace('/', '-')
                    
                except Exception as e:
                    st.error(f"Erro ao processar a requisição: {e}")

        # EXIBIÇÃO, DOWNLOAD E RECOMEÇAR
        if "texto_gerado" in st.session_state:
            st.success("Prontuário Gerado com Sucesso!")
            st.markdown("---")
            
            texto_original = st.session_state["texto_gerado"]
            nome_arq = st.session_state["nome_arquivo"]
            data_arq = st.session_state["data_arquivo"]
            
            # --- 1. EXIBIÇÃO DO TEXTO FORMATADO ---
            texto_formatado_tela = texto_original.replace("\n", "\n\n")
            st.markdown(texto_formatado_tela)
            st.markdown("---")
            
            # --- 2. BOTÕES DE DOWNLOAD LADO A LADO ---
            col1, col2 = st.columns(2)
            
            with col1:
                st.download_button(
                    label="📥 Baixar em Word (.DOCX)",
                    data=gerar_docx(texto_original),
                    file_name=f"prontuario_{nome_arq}_{data_arq}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
            with col2:
                st.download_button(
                    label="📄 Baixar em Texto (.TXT)",
                    data=texto_original,
                    file_name=f"prontuario_{nome_arq}_{data_arq}.txt",
                    mime="text/plain",
                    use_container_width=True
                )
            
            st.write("") # Espaçamento fino
            
            # --- 3. BOTÃO DE RESET NO FINAL ---
            if st.button("🔄 Nova Consulta / Recomeçar", use_container_width=True):
                for key in list(st.session_state.keys()):
                    del st.session_state[key]
                st.rerun()
    else:
        if "texto_gerado" in st.session_state:
            for key in list(st.session_state.keys()):
                del st.session_state[key]